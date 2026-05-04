"""
Deal memo generator.

Produces a professional deal analysis memo in either Word (.docx) or PDF
format from the output of a completed analysis. No new LLM calls — the
memo is assembled purely from data already in the analysis result.

Style: WSJ/Economist-adjacent editorial. Serif display type, black ink,
restrained rules, dense but readable tables.
"""

import io
from datetime import datetime

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
    HRFlowable,
    PageBreak,
    KeepTogether,
)


# -----------------------------------------------------------------------------
# Formatting helpers shared between Word and PDF
# -----------------------------------------------------------------------------

RATIO_ORDER = [
    ("Liquidity", "Current Ratio", "liquidity", "current_ratio", "x"),
    ("Liquidity", "Quick Ratio", "liquidity", "quick_ratio", "x"),
    ("Profitability", "Gross Margin", "profitability", "gross_margin", "%"),
    ("Profitability", "Operating Margin", "profitability", "operating_margin", "%"),
    ("Profitability", "Net Margin", "profitability", "net_margin", "%"),
    ("Profitability", "Return on Assets", "profitability", "return_on_assets", "%"),
    ("Profitability", "Return on Equity", "profitability", "return_on_equity", "%"),
    ("Leverage", "Debt-to-Equity", "leverage", "debt_to_equity", "x"),
    ("Leverage", "Debt-to-Assets", "leverage", "debt_to_assets", "x"),
    ("Leverage", "Interest Coverage", "leverage", "interest_coverage", "x"),
    ("Efficiency", "Asset Turnover", "efficiency", "asset_turnover", "x"),
]


def format_ratio(value, fmt):
    """Format a ratio value. Returns 'N/A' for None."""
    if value is None:
        return "N/A"
    if fmt == "%":
        return f"{value * 100:.2f}%"
    return f"{value:.2f}x"


def format_pct(value):
    """Format a trend percentage change with sign."""
    if value is None:
        return "N/A"
    sign = "+" if value >= 0 else ""
    return f"{sign}{value * 100:.1f}%"


# =============================================================================
# WORD (.docx) GENERATION
# =============================================================================


def _set_cell_background(cell, color_hex):
    """Apply a background fill color to a Word table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _add_bottom_border(paragraph, color_hex="000000", size=8):
    """Add a bottom border rule under a paragraph (used as a horizontal rule)."""
    p_pr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    p_pr.append(pBdr)


def generate_docx(analysis: dict) -> bytes:
    """Build a deal memo as a Word document. Returns raw .docx bytes."""
    extracted = analysis.get("extracted_data", {}) or {}
    narrative = analysis.get("narrative", {}) or {}
    periods = extracted.get("periods", []) or []
    all_ratios = analysis.get("calculated_ratios", []) or []
    trends = analysis.get("trends", {}) or {}

    company = extracted.get("company_name") or "Subject Target"
    currency = extracted.get("currency") or "USD"

    doc = Document()

    # Page setup — US Letter, 1 inch margins
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base body font
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(10.5)

    # ---- Letterhead / masthead ----
    masthead = doc.add_paragraph()
    masthead.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = masthead.add_run("LEDGER")
    run.font.name = "Georgia"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)

    sub = doc.add_paragraph()
    sub_run = sub.add_run("DEAL ANALYSIS MEMO")
    sub_run.font.name = "Georgia"
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    _add_bottom_border(sub, color_hex="000000", size=8)

    # ---- Metadata block ----
    doc.add_paragraph()
    meta_tbl = doc.add_table(rows=3, cols=2)
    meta_tbl.autofit = False
    meta_data = [
        ("TARGET", company.upper()),
        ("PERIODS COVERED", ", ".join(p.get("period_label", "") for p in periods)),
        ("DATE PREPARED", datetime.now().strftime("%B %d, %Y")),
    ]
    for i, (label, value) in enumerate(meta_data):
        label_cell = meta_tbl.cell(i, 0)
        value_cell = meta_tbl.cell(i, 1)
        label_cell.width = Inches(2.0)
        value_cell.width = Inches(4.5)

        lp = label_cell.paragraphs[0]
        lr = lp.add_run(label)
        lr.font.name = "Georgia"
        lr.font.size = Pt(8)
        lr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        lr.bold = True

        vp = value_cell.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.name = "Georgia"
        vr.font.size = Pt(10)
        vr.bold = True

    # Remove borders on metadata table
    for row in meta_tbl.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                border = OxmlElement(f"w:{edge}")
                border.set(qn("w:val"), "nil")
                tc_borders.append(border)
            tc_pr.append(tc_borders)

    doc.add_paragraph()

    # ---- Executive Summary ----
    _add_section_heading(doc, "EXECUTIVE SUMMARY")
    p = doc.add_paragraph(narrative.get("summary", "—"))
    p.paragraph_format.space_after = Pt(10)

    # ---- Period-over-Period Trends ----
    if trends and "revenue_growth" in trends:
        _add_section_heading(doc, "PERIOD-OVER-PERIOD TRENDS")
        trend_tbl = doc.add_table(rows=2, cols=3)
        trend_tbl.style = "Light Grid Accent 1"
        trend_headers = ["Revenue Growth", "Net Income Growth", "Operating Income Growth"]
        trend_values = [
            format_pct(trends.get("revenue_growth")),
            format_pct(trends.get("net_income_growth")),
            format_pct(trends.get("operating_income_growth")),
        ]
        for i, header in enumerate(trend_headers):
            cell = trend_tbl.cell(0, i)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(header)
            run.font.name = "Georgia"
            run.font.size = Pt(8)
            run.bold = True
            _set_cell_background(cell, "EFEBE2")
        for i, value in enumerate(trend_values):
            cell = trend_tbl.cell(1, i)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(value)
            run.font.name = "Georgia"
            run.font.size = Pt(14)
            run.bold = True

        doc.add_paragraph()
        commentary = narrative.get("trend_commentary", "")
        if commentary and commentary != "Insufficient periods for trend analysis.":
            cp = doc.add_paragraph(commentary)
            cp.paragraph_format.left_indent = Inches(0.25)
            cp.paragraph_format.space_after = Pt(10)
            for run in cp.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    # ---- Financial Ratios table ----
    _add_section_heading(doc, "FINANCIAL RATIOS")
    note = doc.add_paragraph()
    note_run = note.add_run("All ratios calculated in deterministic Python. No LLM-derived figures.")
    note_run.italic = True
    note_run.font.size = Pt(8)
    note_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    n_periods = len(periods)
    ratio_tbl = doc.add_table(rows=len(RATIO_ORDER) + 1, cols=2 + n_periods)

    # Header row
    headers = ["Category", "Ratio"] + [p.get("period_label", "") for p in periods]
    for i, h in enumerate(headers):
        cell = ratio_tbl.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.font.name = "Georgia"
        run.font.size = Pt(9)
        run.bold = True
        _set_cell_background(cell, "0A0A0A")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if i >= 2:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Data rows
    for row_idx, (cat, name, group, key, fmt) in enumerate(RATIO_ORDER, start=1):
        cat_cell = ratio_tbl.cell(row_idx, 0)
        cat_run = cat_cell.paragraphs[0].add_run(cat.upper())
        cat_run.font.name = "Georgia"
        cat_run.font.size = Pt(7.5)
        cat_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        name_cell = ratio_tbl.cell(row_idx, 1)
        name_run = name_cell.paragraphs[0].add_run(name)
        name_run.font.name = "Georgia"
        name_run.font.size = Pt(9.5)

        for p_idx, period_ratios in enumerate(all_ratios):
            value = period_ratios.get(group, {}).get(key)
            cell = ratio_tbl.cell(row_idx, 2 + p_idx)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = cell.paragraphs[0].add_run(format_ratio(value, fmt))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)

        # Zebra striping
        if row_idx % 2 == 0:
            for c_idx in range(2 + n_periods):
                _set_cell_background(ratio_tbl.cell(row_idx, c_idx), "F7F5F0")

    doc.add_paragraph()

    # ---- Strengths ----
    _add_section_heading(doc, "INVESTMENT HIGHLIGHTS")
    for item in narrative.get("strengths", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    # ---- Concerns ----
    _add_section_heading(doc, "AREAS OF CONCERN")
    for item in narrative.get("concerns", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    # ---- Key Risks ----
    _add_section_heading(doc, "KEY RISKS")
    risks = narrative.get("red_flags", []) or []
    if risks:
        for item in risks:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item)
    else:
        p = doc.add_paragraph("No material risks identified at this stage of analysis.")
        for run in p.runs:
            run.italic = True
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # ---- Footer ----
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    _add_bottom_border(footer_para, color_hex="CCCCCC", size=4)
    disclaimer = doc.add_paragraph()
    dis_run = disclaimer.add_run(
        "This memo was generated by Ledger using the Gemini API. "
        "All financial ratios are calculated in Python from extracted source data. "
        "Narrative analysis should be reviewed by a qualified analyst before any investment decision."
    )
    dis_run.font.size = Pt(7.5)
    dis_run.italic = True
    dis_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Return as bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_section_heading(doc, text):
    """Add a black-bar section heading to a Word document."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Georgia"
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)
    _add_bottom_border(p, color_hex="0A0A0A", size=12)


# =============================================================================
# PDF GENERATION (ReportLab)
# =============================================================================


def generate_pdf(analysis: dict) -> bytes:
    """Build a deal memo as a PDF. Returns raw PDF bytes."""
    extracted = analysis.get("extracted_data", {}) or {}
    narrative = analysis.get("narrative", {}) or {}
    periods = extracted.get("periods", []) or []
    all_ratios = analysis.get("calculated_ratios", []) or []
    trends = analysis.get("trends", {}) or {}

    company = extracted.get("company_name") or "Subject Target"

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=0.8 * inch,
        rightMargin=0.8 * inch,
        topMargin=0.7 * inch,
        bottomMargin=0.7 * inch,
        title=f"Deal Analysis Memo — {company}",
    )

    # Style registry
    styles = getSampleStyleSheet()
    masthead_style = ParagraphStyle(
        "Masthead",
        parent=styles["Normal"],
        fontName="Times-Bold",
        fontSize=24,
        textColor=colors.black,
        leading=26,
        spaceAfter=2,
    )
    subhead_style = ParagraphStyle(
        "Subhead",
        parent=styles["Normal"],
        fontName="Times-Roman",
        fontSize=9,
        textColor=colors.HexColor("#606060"),
        leading=12,
        spaceAfter=0,
    )
    section_style = ParagraphStyle(
        "Section",
        parent=styles["Normal"],
        fontName="Times-Bold",
        fontSize=11,
        textColor=colors.black,
        leading=14,
        spaceBefore=14,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName="Times-Roman",
        fontSize=10.5,
        textColor=colors.black,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=8,
    )
    bullet_style = ParagraphStyle(
        "Bullet",
        parent=body_style,
        leftIndent=18,
        bulletIndent=6,
        spaceAfter=4,
    )
    note_style = ParagraphStyle(
        "Note",
        parent=styles["Normal"],
        fontName="Times-Italic",
        fontSize=8,
        textColor=colors.HexColor("#808080"),
        leading=10,
        spaceAfter=6,
    )
    meta_label_style = ParagraphStyle(
        "MetaLabel",
        parent=styles["Normal"],
        fontName="Times-Bold",
        fontSize=8,
        textColor=colors.HexColor("#606060"),
        leading=10,
    )
    meta_value_style = ParagraphStyle(
        "MetaValue",
        parent=styles["Normal"],
        fontName="Times-Bold",
        fontSize=10,
        textColor=colors.black,
        leading=12,
    )

    story = []

    # Masthead
    story.append(Paragraph("LEDGER", masthead_style))
    story.append(Paragraph("DEAL ANALYSIS MEMO", subhead_style))
    story.append(Spacer(1, 2))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.black, spaceAfter=12))

    # Metadata block
    meta_data = [
        [Paragraph("TARGET", meta_label_style), Paragraph(company.upper(), meta_value_style)],
        [Paragraph("PERIODS COVERED", meta_label_style),
         Paragraph(", ".join(p.get("period_label", "") for p in periods), meta_value_style)],
        [Paragraph("DATE PREPARED", meta_label_style),
         Paragraph(datetime.now().strftime("%B %d, %Y"), meta_value_style)],
    ]
    meta_tbl = Table(meta_data, colWidths=[1.8 * inch, 5.0 * inch])
    meta_tbl.setStyle(
        TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ])
    )
    story.append(meta_tbl)
    story.append(Spacer(1, 10))

    # Executive Summary
    story.append(_section("EXECUTIVE SUMMARY", section_style))
    story.append(Paragraph(narrative.get("summary", "—"), body_style))

    # Trends
    if trends and "revenue_growth" in trends:
        story.append(_section("PERIOD-OVER-PERIOD TRENDS", section_style))
        trend_data = [
            ["Revenue Growth", "Net Income Growth", "Operating Income Growth"],
            [
                format_pct(trends.get("revenue_growth")),
                format_pct(trends.get("net_income_growth")),
                format_pct(trends.get("operating_income_growth")),
            ],
        ]
        trend_tbl = Table(trend_data, colWidths=[2.2 * inch, 2.3 * inch, 2.3 * inch])
        trend_tbl.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 8),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EFEBE2")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#606060")),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("FONTNAME", (0, 1), (-1, 1), "Times-Bold"),
            ("FONTSIZE", (0, 1), (-1, 1), 15),
            ("TOPPADDING", (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
        ]))
        story.append(trend_tbl)

        commentary = narrative.get("trend_commentary", "")
        if commentary and commentary != "Insufficient periods for trend analysis.":
            story.append(Spacer(1, 6))
            comment_style = ParagraphStyle(
                "Commentary",
                parent=body_style,
                fontName="Times-Italic",
                fontSize=9.5,
                textColor=colors.HexColor("#404040"),
                leftIndent=10,
            )
            story.append(Paragraph(commentary, comment_style))

    # Financial Ratios
    story.append(_section("FINANCIAL RATIOS", section_style))
    story.append(Paragraph(
        "All ratios calculated in deterministic Python. No LLM-derived figures.",
        note_style,
    ))

    # Build the ratios table
    period_labels = [p.get("period_label", "") for p in periods]
    ratio_header = ["Category", "Ratio"] + period_labels
    ratio_rows = [ratio_header]
    for cat, name, group, key, fmt in RATIO_ORDER:
        row = [cat.upper(), name]
        for period_ratios in all_ratios:
            value = period_ratios.get(group, {}).get(key)
            row.append(format_ratio(value, fmt))
        ratio_rows.append(row)

    # Dynamic column widths
    n_periods = len(periods)
    total_width = 6.9 * inch
    cat_w = 1.1 * inch
    name_w = 1.9 * inch
    period_w = (total_width - cat_w - name_w) / max(n_periods, 1)
    col_widths = [cat_w, name_w] + [period_w] * n_periods

    ratio_tbl = Table(ratio_rows, colWidths=col_widths, repeatRows=1)
    table_style_cmds = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0A0A0A")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 9),
        ("ALIGN", (0, 0), (1, -1), "LEFT"),
        ("ALIGN", (2, 0), (-1, -1), "RIGHT"),
        ("FONTNAME", (0, 1), (0, -1), "Times-Roman"),
        ("FONTSIZE", (0, 1), (0, -1), 7.5),
        ("TEXTCOLOR", (0, 1), (0, -1), colors.HexColor("#808080")),
        ("FONTNAME", (1, 1), (1, -1), "Times-Roman"),
        ("FONTSIZE", (1, 1), (1, -1), 9.5),
        ("FONTNAME", (2, 1), (-1, -1), "Courier"),
        ("FONTSIZE", (2, 1), (-1, -1), 9),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LINEBELOW", (0, 0), (-1, 0), 0.5, colors.black),
        ("LINEBELOW", (0, -1), (-1, -1), 0.5, colors.black),
    ]
    # Zebra stripes
    for r in range(2, len(ratio_rows), 2):
        table_style_cmds.append(("BACKGROUND", (0, r), (-1, r), colors.HexColor("#F7F5F0")))
    ratio_tbl.setStyle(TableStyle(table_style_cmds))
    story.append(ratio_tbl)

    # Investment Highlights
    story.append(_section("INVESTMENT HIGHLIGHTS", section_style))
    for item in narrative.get("strengths", []) or []:
        story.append(Paragraph(f"• {item}", bullet_style))

    # Areas of Concern
    story.append(_section("AREAS OF CONCERN", section_style))
    for item in narrative.get("concerns", []) or []:
        story.append(Paragraph(f"• {item}", bullet_style))

    # Key Risks
    story.append(_section("KEY RISKS", section_style))
    risks = narrative.get("red_flags", []) or []
    if risks:
        for item in risks:
            story.append(Paragraph(f"• {item}", bullet_style))
    else:
        story.append(Paragraph(
            "No material risks identified at this stage of analysis.",
            ParagraphStyle(
                "NoRisks",
                parent=body_style,
                fontName="Times-Italic",
                textColor=colors.HexColor("#606060"),
            ),
        ))

    # Footer disclaimer
    story.append(Spacer(1, 16))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#CCCCCC")))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "This memo was generated by Ledger using the Gemini API. All financial ratios "
        "are calculated in Python from extracted source data. Narrative analysis should "
        "be reviewed by a qualified analyst before any investment decision.",
        note_style,
    ))

    doc.build(story)
    return buf.getvalue()


def _section(text, style):
    """Return a Paragraph for a section heading with a bottom rule applied."""
    # We can't easily attach a bottom rule to a Paragraph in ReportLab,
    # so we use a KeepTogether pattern in the caller if needed.
    return Paragraph(text, style)
